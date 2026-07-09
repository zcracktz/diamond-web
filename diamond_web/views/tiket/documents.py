"""Tiket document generation — DOCX downloads for Tanda Terima, Lampiran, and Register Data."""

import re
from io import BytesIO

from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.views.decorators.http import require_GET

from ...models.detil_tanda_terima import DetilTandaTerima
from ...models.klasifikasi_jenis_data import KlasifikasiJenisData
from ...models.tiket import Tiket
from ...models.tiket_pic import TiketPIC
from ...models.docx_template import DocxTemplate
from ...utils.docx_template import fill_template_with_data
from ...utils import format_number_with_separator, format_periode


def _is_p3de_user(user):
    """Check if user is P3DE (can generate/download documents).

    Args:
        user: Django User object to check permissions for.

    Returns:
        bool: True if user is superuser, admin, or in user_p3de group.
    """
    if not user or not user.is_authenticated:
        return False
    if user.is_superuser or user.groups.filter(name='admin').exists():
        return True
    return user.groups.filter(name='user_p3de').exists()


def _format_periode_tiket(tiket_obj):
    """Format periode display for a tiket object using periode penerimaan.
    
    This wrapper extracts periode information from the tiket object and calls
    the centralized format_periode function with Roman numerals for semester,
    triwulan, and kuartal.
    """
    if not tiket_obj.id_periode_data or not tiket_obj.id_periode_data.id_periode_pengiriman:
        return '-'

    periode_desc = tiket_obj.id_periode_data.id_periode_pengiriman.periode_penerimaan or '-'
    tahun = tiket_obj.tahun if tiket_obj.tahun else None
    periode = tiket_obj.periode

    if tahun is None:
        return '-'

    return format_periode(periode_desc, periode, tahun)


def _safe_filename_part(raw):
    """Convert a string into a filename-safe format.

    Replaces all non-alphanumeric characters (except dots, hyphens, underscores)
    with underscores and strips leading/trailing underscores.

    Args:
        raw: Input string to sanitize for use in filenames.

    Returns:
        str: Filename-safe string, or 'file' if result would be empty.
    """
    return re.sub(r'[^A-Za-z0-9._-]+', '_', str(raw or '')).strip('_') or 'file'


def _format_date_indonesian(date_obj):
    """Format a date object as D Bulan YYYY in Indonesian with capitalized month names.

    Args:
        date_obj: datetime.date or datetime.datetime object to format.

    Returns:
        str: Formatted date string (e.g., '15 Januari 2026'), or '-' if date_obj is None.
    """
    if not date_obj:
        return '-'
    
    bulan_map = {
        1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April', 5: 'Mei', 6: 'Juni',
        7: 'Juli', 8: 'Agustus', 9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
    }
    bulan = bulan_map.get(date_obj.month, '')
    return f"{date_obj.day} {bulan} {date_obj.year}"


def _build_table_doc(title, headers, rows_data):
    """Build a DOCX document with a single table.

    Creates a new python-docx Document, adds a heading with the given title,
    and populates a styled table with the provided headers and row data.

    Args:
        title (str): Heading text for the document.
        headers (list of str): Column header labels for the table.
        rows_data (list of list): Row data, where each inner list contains
            string values for each column.

    Returns:
        Document: A python-docx Document object ready for saving or streaming.
    """
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows_data:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
    return doc


@login_required
@user_passes_test(lambda u: _is_p3de_user(u))
@require_GET
def _merge_docx(doc1, doc2):
    """Combine two docx Documents by adding a page break and appending elements."""
    doc1.add_page_break()
    import copy
    for element in list(doc2.element.body):
        # Skip section properties element to avoid page setup bugs
        if element.tag.endswith('sectPr'):
            continue
        doc1.element.body.append(copy.deepcopy(element))
    return doc1


def _generate_single_document(request, pk, doc_type):
    """Inner helper to generate a single docx Document (template or fallback)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError('Library python-docx belum terpasang.')

    tiket = get_object_or_404(
        Tiket.objects.select_related(
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap__id_kategori',
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
            'id_periode_data__id_periode_pengiriman',
            'id_bentuk_data',
            'id_cara_penyampaian',
        ).prefetch_related(
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap__ilap_kpp_relations__id_kpp__id_kanwil',
        ),
        pk=pk,
    )

    if not request.user.groups.filter(name='admin').exists() and not request.user.is_superuser:
        has_access = TiketPIC.objects.filter(id_tiket=tiket, id_user=request.user, active=True).exists()
        if not has_access:
            raise PermissionError('Tidak memiliki akses ke tiket ini.')

    # Collect tanda-terima group and associated tiket rows
    detil = DetilTandaTerima.objects.select_related('id_tanda_terima').filter(id_tiket=tiket).order_by('-id').first()
    tanda_terima = detil.id_tanda_terima if detil else None

    if tanda_terima:
        tiket_ids = list(
            DetilTandaTerima.objects.filter(id_tanda_terima=tanda_terima).values_list('id_tiket_id', flat=True)
        )
        tiket_rows = list(
            Tiket.objects.filter(id__in=tiket_ids).select_related(
                'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                'id_periode_data__id_periode_pengiriman',
                'id_bentuk_data',
                'id_cara_penyampaian',
            ).order_by('id')
        )
    else:
        tiket_rows = [tiket]

    # Dasar hukum lookup
    jenis_data_ids = {
        t.id_periode_data.id_sub_jenis_data_ilap_id
        for t in tiket_rows
        if t.id_periode_data and t.id_periode_data.id_sub_jenis_data_ilap_id
    }
    dasar_hukum_map = {}
    for row in KlasifikasiJenisData.objects.filter(id_sub_jenis_data_id__in=jenis_data_ids).select_related('id_klasifikasi_tabel'):
        dasar_hukum_map.setdefault(row.id_sub_jenis_data_id, []).append(row.id_klasifikasi_tabel.deskripsi)

    # PIC P3DE name
    p3de = TiketPIC.objects.select_related('id_user').filter(
        id_tiket=tiket,
        role=TiketPIC.Role.P3DE,
        active=True,
    ).order_by('id').first()
    p3de_name = '-'
    if p3de and p3de.id_user:
        p3de_name = p3de.id_user.get_full_name().strip() or p3de.id_user.username

    # Derived fields
    ilap = (
        tiket.id_periode_data.id_sub_jenis_data_ilap.id_ilap
        if tiket.id_periode_data and tiket.id_periode_data.id_sub_jenis_data_ilap
        else None
    )
    # For regional ILAPs (those with KPP), use nama_kanwil; otherwise use nama_ilap
    if ilap:
        first_kpp_rel = ilap.ilap_kpp_relations.select_related('id_kpp__id_kanwil').first()
        if first_kpp_rel and first_kpp_rel.id_kpp and first_kpp_rel.id_kpp.id_kanwil:
            diterima_dari = first_kpp_rel.id_kpp.id_kanwil.nama_kanwil
        else:
            diterima_dari = ilap.nama_ilap
    else:
        diterima_dari = '-'

    # Collect multi-value fields from tiket_rows (deduplicated)
    periode_list, nomor_surat_list, tanggal_surat_list = [], [], []
    bentuk_data_list, cara_penyampaian_list = [], []
    
    seen_periode, seen_nomor_surat, seen_tanggal_surat = set(), set(), set()
    seen_bentuk_data, seen_cara_penyampaian = set(), set()
    
    for t in tiket_rows:
        label = _format_periode_tiket(t)
        if label not in seen_periode:
            seen_periode.add(label)
            periode_list.append(label)
        
        nomor_surat = t.nomor_surat_pengantar or '-'
        if nomor_surat not in seen_nomor_surat:
            seen_nomor_surat.add(nomor_surat)
            nomor_surat_list.append(nomor_surat)
        
        tanggal = _format_date_indonesian(t.tanggal_surat_pengantar) if t.tanggal_surat_pengantar else '-'
        if tanggal not in seen_tanggal_surat:
            seen_tanggal_surat.add(tanggal)
            tanggal_surat_list.append(tanggal)
        
        bentuk = t.id_bentuk_data.deskripsi if t.id_bentuk_data else '-'
        if bentuk not in seen_bentuk_data:
            seen_bentuk_data.add(bentuk)
            bentuk_data_list.append(bentuk)
        
        cara = t.id_cara_penyampaian.deskripsi if t.id_cara_penyampaian else '-'
        if cara not in seen_cara_penyampaian:
            seen_cara_penyampaian.add(cara)
            cara_penyampaian_list.append(cara)

    nomor_tanda_terima = tanda_terima.nomor_tanda_terima_format if tanda_terima else '-'
    tgl_terima_dip = _format_date_indonesian(tiket.tgl_terima_dip) if tiket.tgl_terima_dip else '-'

    tahun_data_list = sorted({str(t.tahun) for t in tiket_rows if t.tahun})

    # Build variable dictionary for template filling
    template_variables = {
        '{{nomor_tiket}}': nomor_tanda_terima,
        '{{nomor_tanda_terima}}': nomor_tanda_terima,
        '{{tanggal_tanda_terima}}': _format_date_indonesian(tanda_terima.tanggal_tanda_terima) if tanda_terima else '-',
        '{{tahun_data}}': ', '.join(tahun_data_list) if tahun_data_list else '-',
        '{{diterima_dari}}': diterima_dari,
        '{{nama_kantor}}': diterima_dari,
        '{{nomor_surat_pengantar}}': ', '.join(nomor_surat_list) if nomor_surat_list else '-',
        '{{tanggal_surat_pengantar}}': ', '.join(tanggal_surat_list) if tanggal_surat_list else '-',
        '{{tanggal_penerimaan}}': ', '.join(tanggal_surat_list) if tanggal_surat_list else '-',
        '{{nama_ilap}}': ilap.nama_ilap if ilap else '-',
        '{{jenis_data}}': 'Terlampir',
        '{{periode_data}}': ', '.join(periode_list) if periode_list else '-',
        '{{bentuk_data}}': ', '.join(bentuk_data_list) if bentuk_data_list else '-',
        '{{tanggal_terima_dip}}': tgl_terima_dip,
        '{{cara_penyampaian}}': ', '.join(cara_penyampaian_list) if cara_penyampaian_list else '-',
        '{{nama_pic_p3de}}': p3de_name,
        '{{nama_pic}}': p3de_name,
        '{{email_pic}}': '-',
        '{{telepon_pic}}': '-',
        '{{nama_tabel}}': 'Terlampir',
        '{{jumlah_record}}': '-',
        '{{ukuran_file}}': '-',
    }

    # DOC Type Selection and Template Processing
    now_ts = timezone.now().strftime('%Y%m%d_%H%M%S_%f')
    nomor_safe = _safe_filename_part(nomor_tanda_terima)

    # Determine region type (Regional or Nasional/Internasional)
    region_type = 'regional'
    if ilap and ilap.id_kategori_wilayah:
        kategori_wilayah_desc = ilap.id_kategori_wilayah.deskripsi.lower()
        if 'regional' not in kategori_wilayah_desc:
            region_type = 'nasional_internasional'
    
    # Map doc_type to template jenis_dokumen based on region type
    doc_type_map = {
        'tanda_terima_only': f'tanda_terima_{region_type}',
        'tanda_terima': f'tanda_terima_{region_type}',
        'lampiran': f'lampiran_tanda_terima_{region_type}',
        'register': 'register_penerimaan_data',
        'pkdi_lengkap': f'surat_pkdi_{region_type}_lengkap',
        'pkdi_sebagian': f'surat_pkdi_{region_type}_sebagian',
        'klarifikasi': 'surat_klarifikasi',
        'nd_pengantar': 'nd_pengantar_pide',
    }
    template_jenis = doc_type_map.get(doc_type, f'tanda_terima_{region_type}')

    # Try to find an active template for this document type
    template = DocxTemplate.objects.filter(
        jenis_dokumen=template_jenis,
        active=True
    ).first()

    if template and template.file_template:
        try:
            row_data = None
            if doc_type in ('lampiran', 'register'):
                row_data = []
                nomor_counter = 1
                for t in tiket_rows:
                    sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
                    ilap_obj = sub.id_ilap if sub else None
                    dasar_hukum_list = dasar_hukum_map.get(sub.id, []) if sub else []
                    
                    nama_kanwil = '-'
                    if ilap_obj:
                        first_kpp_rel = ilap_obj.ilap_kpp_relations.select_related('id_kpp__id_kanwil').first()
                        if first_kpp_rel and first_kpp_rel.id_kpp and first_kpp_rel.id_kpp.id_kanwil:
                            nama_kanwil = first_kpp_rel.id_kpp.id_kanwil.nama_kanwil
                    
                    status_data = '-'
                    if sub and sub.id_status_data:
                        status_data = sub.id_status_data.deskripsi
                    
                    if doc_type == 'lampiran':
                        row_data.append({
                            'nomor': str(nomor_counter),
                            'nama_kanwil': nama_kanwil,
                            'nama_ilap': ilap_obj.nama_ilap if ilap_obj else '-',
                            'sub_jenis_data': sub.nama_sub_jenis_data if sub else '-',
                            'periode_data': _format_periode_tiket(t),
                            'status_data': status_data,
                            'jumlah_baris_diterima': format_number_with_separator(t.baris_diterima),
                            'dasar_hukum': ', '.join(dasar_hukum_list) if dasar_hukum_list else '-',
                        })
                    else:  # register
                        row_data.append({
                            'nomor': str(nomor_counter),
                            'nama_kanwil': nama_kanwil,
                            'nama_ilap': ilap_obj.nama_ilap if ilap_obj else '-',
                            'sub_jenis_data': sub.nama_sub_jenis_data if sub else '-',
                            'periode_data': _format_periode_tiket(t),
                            'status_data': status_data,
                            'jumlah_baris_diterima': format_number_with_separator(t.baris_diterima),
                            'dasar_hukum': ', '.join(dasar_hukum_list) if dasar_hukum_list else '-',
                        })
                    nomor_counter += 1
            elif doc_type in ('pkdi_lengkap', 'pkdi_sebagian', 'klarifikasi'):
                row_data = []
                nomor_counter = 1
                for t in tiket_rows:
                    sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
                    if doc_type == 'klarifikasi':
                        row_data.append({
                            'nomor': str(nomor_counter),
                            'sub_jenis_data': sub.nama_sub_jenis_data if sub else '-',
                            'periode_data': _format_periode_tiket(t),
                            'jumlah_baris_diterima': format_number_with_separator(t.baris_diterima) if t.baris_diterima is not None else '-',
                            'jumlah_baris_tidak_lengkap': format_number_with_separator(t.baris_tidak_lengkap) if t.baris_tidak_lengkap is not None else '-',
                        })
                    else:
                        row_data.append({
                            'nomor': str(nomor_counter),
                            'sub_jenis_data': sub.nama_sub_jenis_data if sub else '-',
                            'periode_data': _format_periode_tiket(t),
                            'jumlah_baris_diterima': format_number_with_separator(t.baris_diterima) if t.baris_diterima is not None else '-',
                            'jumlah_baris_lengkap': format_number_with_separator(t.baris_lengkap) if t.baris_lengkap is not None else '-',
                            'jumlah_baris_tidak_lengkap': format_number_with_separator(t.baris_tidak_lengkap) if t.baris_tidak_lengkap is not None else '-',
                        })
                    nomor_counter += 1

            doc_buffer = fill_template_with_data(template.file_template.open('rb'), template_variables, row_data=row_data)
            
            if doc_type == 'lampiran':
                filename = f'lampiran_tanda_terima_{nomor_safe}_{now_ts}.docx'
            elif doc_type == 'register':
                filename = f'register_data_{nomor_safe}_{now_ts}.docx'
            elif doc_type == 'pkdi_lengkap':
                filename = f'surat_pkdi_lengkap_{now_ts}.docx'
            elif doc_type == 'pkdi_sebagian':
                filename = f'surat_pkdi_lengkap_sebagian_{now_ts}.docx'
            elif doc_type == 'klarifikasi':
                filename = f'surat_klarifikasi_{now_ts}.docx'
            elif doc_type == 'nd_pengantar':
                filename = f'nd_pengantar_pide_{nomor_safe}_{now_ts}.docx'
            else:
                filename = f'tanda_terima_{nomor_safe}_{now_ts}.docx'
            
            return Document(doc_buffer), filename
        except Exception:
            pass

    # Fallback: Generate default documents
    if doc_type == 'lampiran':
        lampiran_headers = ['Nama ILAP', 'Jenis Data', 'Periode Data Tahun', 'Baris Diterima', 'Dasar Hukum']
        lampiran_rows = []
        for t in tiket_rows:
            sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
            ilap_obj = sub.id_ilap if sub else None
            dasar_hukum_list = dasar_hukum_map.get(sub.id, []) if sub else []
            lampiran_rows.append([
                f"{ilap_obj.id_ilap} - {ilap_obj.nama_ilap}" if ilap_obj else '-',
                f"{sub.id_sub_jenis_data} - {sub.nama_sub_jenis_data}" if sub else '-',
                _format_periode_tiket(t),
                str(t.baris_diterima if t.baris_diterima is not None else '-'),
                ', '.join(dasar_hukum_list) if dasar_hukum_list else '-',
            ])
        doc = _build_table_doc('Lampiran Tanda Terima', lampiran_headers, lampiran_rows)
        filename = f'lampiran_tanda_terima_{nomor_safe}_{now_ts}.docx'
    elif doc_type == 'register':
        register_headers = ['Nama ILAP', 'Jenis Data', 'Periode Data Tahun', 'Baris Diterima', 'Dasar Hukum']
        register_rows = []
        for t in tiket_rows:
            sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
            ilap_obj = sub.id_ilap if sub else None
            dasar_hukum_list = dasar_hukum_map.get(sub.id, []) if sub else []
            register_rows.append([
                f"{ilap_obj.id_ilap} - {ilap_obj.nama_ilap}" if ilap_obj else '-',
                f"{sub.id_sub_jenis_data} - {sub.nama_sub_jenis_data}" if sub else '-',
                _format_periode_tiket(t),
                str(t.baris_diterima if t.baris_diterima is not None else '-'),
                ', '.join(dasar_hukum_list) if dasar_hukum_list else '-',
            ])
        doc = _build_table_doc('Register Data', register_headers, register_rows)
        filename = f'register_data_{nomor_safe}_{now_ts}.docx'
    elif doc_type == 'pkdi_lengkap':
        doc_pkdi = Document()
        doc_pkdi.add_heading('Surat PKDI Lengkap (Pernyataan Kesesuaian Data)', level=1)
        fields = [
            ('Nomor Tanda Terima', nomor_tanda_terima),
            ('Diterima Dari', diterima_dari),
            ('Nama ILAP', ilap.nama_ilap if ilap else '-'),
            ('Jenis Data', 'Terlampir'),
            ('Periode Data', ', '.join(periode_list) if periode_list else '-'),
            ('Status Data', 'Lengkap - Semua variabel dan satuan data sesuai spesifikasi'),
            ('Tanggal Terima DIP', tgl_terima_dip),
        ]
        table_fields = doc_pkdi.add_table(rows=0, cols=2)
        table_fields.style = 'Table Grid'
        for key, value in fields:
            row = table_fields.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        doc = doc_pkdi
        filename = f'surat_pkdi_lengkap_{now_ts}.docx'
    elif doc_type == 'pkdi_sebagian':
        doc_pkdi = Document()
        doc_pkdi.add_heading('Surat PKDI Lengkap Sebagian (Pernyataan Kesesuaian Data)', level=1)
        fields = [
            ('Nomor Tanda Terima', nomor_tanda_terima),
            ('Diterima Dari', diterima_dari),
            ('Nama ILAP', ilap.nama_ilap if ilap else '-'),
            ('Jenis Data', 'Terlampir'),
            ('Periode Data', ', '.join(periode_list) if periode_list else '-'),
            ('Status Data', 'Lengkap Sebagian - Sebagian variabel dan satuan data sesuai spesifikasi'),
            ('Keterangan', 'Beberapa item masih perlu perbaikan dan akan dikomunikasikan lebih lanjut'),
            ('Tanggal Terima DIP', tgl_terima_dip),
        ]
        table_fields = doc_pkdi.add_table(rows=0, cols=2)
        table_fields.style = 'Table Grid'
        for key, value in fields:
            row = table_fields.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        doc = doc_pkdi
        filename = f'surat_pkdi_lengkap_sebagian_{now_ts}.docx'
    elif doc_type == 'klarifikasi':
        doc_klr = Document()
        doc_klr.add_heading('Surat Klarifikasi Data', level=1)
        fields = [
            ('Nomor Surat', ', '.join(nomor_surat_list) if nomor_surat_list else '-'),
            ('Tanggal', ', '.join(tanggal_surat_list) if tanggal_surat_list else '-'),
            ('Kepada', diterima_dari),
            ('Nama ILAP', ilap.nama_ilap if ilap else '-'),
            ('Nomor Tanda Terima', nomor_tanda_terima),
            ('Status Data', 'Perlu Klarifikasi - Data belum lengkap'),
            ('Petugas P3DE', p3de_name),
            ('Keterangan', 'Mohon dapat dikonfirmasi dan perbaiki sesuai spesifikasi yang diminta'),
        ]
        table_fields = doc_klr.add_table(rows=0, cols=2)
        table_fields.style = 'Table Grid'
        for key, value in fields:
            row = table_fields.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        doc = doc_klr
        filename = f'surat_klarifikasi_{now_ts}.docx'
    else:
        doc_tanda = Document()
        doc_tanda.add_heading('Tanda Terima Data', level=1)
        fields = [
            ('Nomor Tanda Terima',    nomor_tanda_terima),
            ('Diterima Dari',         diterima_dari),
            ('Nomor Surat Pengantar',  ', '.join(nomor_surat_list) if nomor_surat_list else '-'),
            ('Tanggal Surat Pengantar', ', '.join(tanggal_surat_list) if tanggal_surat_list else '-'),
            ('Nama ILAP',             ilap.nama_ilap if ilap else '-'),
            ('Jenis Data',            'Terlampir'),
            ('Periode Data',          ', '.join(periode_list) if periode_list else '-'),
            ('Bentuk Data',           ', '.join(bentuk_data_list) if bentuk_data_list else '-'),
            ('Tanggal Terima DIP',    tgl_terima_dip),
            ('Cara Penyampaian',      ', '.join(cara_penyampaian_list) if cara_penyampaian_list else '-'),
            ('Nama PIC P3DE',         p3de_name),
        ]
        table_fields = doc_tanda.add_table(rows=0, cols=2)
        table_fields.style = 'Table Grid'
        for key, value in fields:
            row = table_fields.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        doc = doc_tanda
        filename = f'tanda_terima_{nomor_safe}_{now_ts}.docx'

    return doc, filename


def tiket_documents_download(request, pk):
    """Generate and download DOCX documents for a tiket.
    
    If doc_type is 'tanda_terima', automatically combines the Tanda Terima
    and its associated Lampiran into a single document with a page break.
    Otherwise, returns the single requested document.
    """
    doc_type = request.GET.get('doc_type', 'tanda_terima')
    
    if doc_type == 'tanda_terima':
        try:
            # Generate Tanda Terima (only) and Lampiran, then merge them
            doc_tanda, filename_tanda = _generate_single_document(request, pk, 'tanda_terima_only')
            doc_lampiran, _ = _generate_single_document(request, pk, 'lampiran')
            
            merged_doc = _merge_docx(doc_tanda, doc_lampiran)
            
            buffer = BytesIO()
            merged_doc.save(buffer)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{filename_tanda}"'
            return response
        except Exception:
            # Fall back to single Tanda Terima if merge fails
            try:
                doc_tanda, filename_tanda = _generate_single_document(request, pk, 'tanda_terima_only')
                buffer = BytesIO()
                doc_tanda.save(buffer)
                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                )
                response['Content-Disposition'] = f'attachment; filename="{filename_tanda}"'
                return response
            except Exception as inner_e:
                return HttpResponse(f'Gagal menghasilkan dokumen: {str(inner_e)}', status=500)
    else:
        # Generate single requested document (register, klarifikasi, etc.)
        try:
            doc, filename = _generate_single_document(request, pk, doc_type)
            buffer = BytesIO()
            doc.save(buffer)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            return HttpResponse(f'Gagal menghasilkan dokumen: {str(e)}', status=500)
