"""Bulk DOCX generation pages for P3DE users.

Page 1:
- Generate PKDI / PKDI Sebagian / Klarifikasi for multiple tickets.

Page 2:
- Generate ND Pengantar PIDE for multiple tickets with status Dikirim ke PIDE.
"""

from datetime import datetime
from io import BytesIO

from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponse
from django.shortcuts import redirect, render
from django.views.decorators.http import require_http_methods

from ..constants.tiket_status import STATUS_DIKIRIM_KE_PIDE
from ..models.detil_tanda_terima import DetilTandaTerima
from ..models.docx_template import DocxTemplate
from ..models.ilap import ILAP
from ..models.klasifikasi_jenis_data import KlasifikasiJenisData
from ..models.tiket import Tiket
from ..utils import format_number_with_separator, format_periode
from ..utils.docx_template import fill_template_with_data
from .mixins import get_active_p3de_ilap_ids


def _is_p3de_user(user):
    if not user or not user.is_authenticated:
        return False
    if user.is_superuser or user.groups.filter(name='admin').exists():
        return True
    return user.groups.filter(name='user_p3de').exists()


def _format_date_indonesian(date_obj):
    if not date_obj:
        return '-'
    bulan_map = {
        1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April', 5: 'Mei', 6: 'Juni',
        7: 'Juli', 8: 'Agustus', 9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember',
    }
    return f"{date_obj.day} {bulan_map.get(date_obj.month, '')} {date_obj.year}"


def _format_periode_tiket(tiket_obj):
    if not tiket_obj.id_periode_data or not tiket_obj.id_periode_data.id_periode_pengiriman:
        return '-'

    periode_desc = tiket_obj.id_periode_data.id_periode_pengiriman.periode_penerimaan or '-'
    tahun = tiket_obj.tahun if tiket_obj.tahun else None
    periode = tiket_obj.periode

    if tahun is None:
        return '-'

    return format_periode(periode_desc, periode, tahun)


def _parse_date(value):
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except (TypeError, ValueError):
        return None


def _base_queryset(ilap_id, tanggal_terima):
    return Tiket.objects.filter(
        id_periode_data__id_sub_jenis_data_ilap__id_ilap_id=ilap_id,
        tgl_terima_dip__date=tanggal_terima,
        tanda_terima=True,
    ).select_related(
        'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
        'id_periode_data__id_periode_pengiriman',
        'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
        'id_status_penelitian',
    ).prefetch_related(
        'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
    ).order_by('id')


def _apply_doc_type_filter(queryset, doc_type):
    if doc_type == 'pkdi_lengkap':
        return queryset.filter(id_status_penelitian__deskripsi='Lengkap')
    if doc_type == 'pkdi_sebagian':
        return queryset.filter(id_status_penelitian__deskripsi='Lengkap Sebagian')
    if doc_type == 'klarifikasi':
        return queryset.filter(id_status_penelitian__deskripsi__in=['Lengkap Sebagian', 'Tidak Lengkap'])
    return queryset.none()


def _build_table_doc(title, headers, rows_data):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows_data:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
    return doc


def _generate_docx_for_tickets(selected_tickets, doc_type, title_prefix):
    """Generate DOCX for selected tickets using template (preferred) or fallback table."""
    if not selected_tickets:
        return None

    first_ticket = selected_tickets[0]
    ilap = first_ticket.id_periode_data.id_sub_jenis_data_ilap.id_ilap

    # Determine region type for PKDI template selection.
    region_type = 'regional'
    if ilap and ilap.id_kategori_wilayah:
        kategori_wilayah_desc = (ilap.id_kategori_wilayah.deskripsi or '').lower()
        if 'regional' not in kategori_wilayah_desc:
            region_type = 'nasional_internasional'

    template_type_map = {
        'tanda_terima': f'tanda_terima_{region_type}',
        'pkdi_lengkap': f'surat_pkdi_{region_type}_lengkap',
        'pkdi_sebagian': f'surat_pkdi_{region_type}_sebagian',
        'klarifikasi': 'surat_klarifikasi',
        'nd_pengantar': 'nd_pengantar_pide',
    }
    template_jenis = template_type_map.get(doc_type)

    # Lookup supporting maps.
    tiket_ids = [t.id for t in selected_tickets]
    detils = DetilTandaTerima.objects.filter(
        id_tiket_id__in=tiket_ids
    ).select_related('id_tanda_terima').order_by('id_tanda_terima__tanggal_tanda_terima')

    tanda_terima_map = {}
    for d in detils:
        if d.id_tiket_id not in tanda_terima_map:
            tanda_terima_map[d.id_tiket_id] = d.id_tanda_terima

    jenis_data_ids = {
        t.id_periode_data.id_sub_jenis_data_ilap_id
        for t in selected_tickets
        if t.id_periode_data and t.id_periode_data.id_sub_jenis_data_ilap_id
    }
    dasar_hukum_map = {}
    for row in KlasifikasiJenisData.objects.filter(
        id_sub_jenis_data_id__in=jenis_data_ids
    ).select_related('id_klasifikasi_tabel'):
        dasar_hukum_map.setdefault(row.id_sub_jenis_data_id, []).append(row.id_klasifikasi_tabel.deskripsi)

    # Header-level variables (aligned with documents.py)
    if ilap and ilap.id_kpp and ilap.id_kpp.id_kanwil:
        diterima_dari = ilap.id_kpp.id_kanwil.nama_kanwil
    else:
        diterima_dari = ilap.nama_ilap if ilap else '-'

    periode_list, nomor_surat_list, tanggal_surat_list = [], [], []
    bentuk_data_list, cara_penyampaian_list = [], []
    nomor_tanda_terima_list, tanggal_tanda_terima_list = [], []
    seen_periode, seen_nomor_surat, seen_tanggal_surat = set(), set(), set()
    seen_bentuk_data, seen_cara_penyampaian = set(), set()
    seen_nomor_tt, seen_tanggal_tt = set(), set()

    row_data = []
    for idx, t in enumerate(selected_tickets, start=1):
        sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
        ilap_obj = sub.id_ilap if sub else None
        status_data = sub.id_status_data.deskripsi if sub and sub.id_status_data else '-'
        status_penelitian = t.id_status_penelitian.deskripsi if t.id_status_penelitian else '-'
        dasar_hukum = ', '.join(dasar_hukum_map.get(sub.id, [])) if sub else '-'

        tt = tanda_terima_map.get(t.id)
        nomor_tt = tt.nomor_tanda_terima_format if tt else '-'
        tanggal_tt = _format_date_indonesian(tt.tanggal_tanda_terima) if tt else '-'

        nama_kanwil = '-'
        if ilap_obj and ilap_obj.id_kpp and ilap_obj.id_kpp.id_kanwil:
            nama_kanwil = ilap_obj.id_kpp.id_kanwil.nama_kanwil

        periode_label = _format_periode_tiket(t)
        if periode_label not in seen_periode:
            seen_periode.add(periode_label)
            periode_list.append(periode_label)

        nomor_surat = t.nomor_surat_pengantar or '-'
        if nomor_surat not in seen_nomor_surat:
            seen_nomor_surat.add(nomor_surat)
            nomor_surat_list.append(nomor_surat)

        tanggal_surat = _format_date_indonesian(t.tanggal_surat_pengantar) if t.tanggal_surat_pengantar else '-'
        if tanggal_surat not in seen_tanggal_surat:
            seen_tanggal_surat.add(tanggal_surat)
            tanggal_surat_list.append(tanggal_surat)

        bentuk_data = t.id_bentuk_data.deskripsi if getattr(t, 'id_bentuk_data', None) else '-'
        if bentuk_data not in seen_bentuk_data:
            seen_bentuk_data.add(bentuk_data)
            bentuk_data_list.append(bentuk_data)

        cara_penyampaian = t.id_cara_penyampaian.deskripsi if getattr(t, 'id_cara_penyampaian', None) else '-'
        if cara_penyampaian not in seen_cara_penyampaian:
            seen_cara_penyampaian.add(cara_penyampaian)
            cara_penyampaian_list.append(cara_penyampaian)

        if nomor_tt not in seen_nomor_tt:
            seen_nomor_tt.add(nomor_tt)
            nomor_tanda_terima_list.append(nomor_tt)

        if tanggal_tt not in seen_tanggal_tt:
            seen_tanggal_tt.add(tanggal_tt)
            tanggal_tanda_terima_list.append(tanggal_tt)

        row_data.append({
            'nomor': str(idx),
            'nomor_tiket': t.nomor_tiket,
            'nomor_tanda_terima': nomor_tt,
            'tanggal_tanda_terima': tanggal_tt,
            'nama_kanwil': nama_kanwil,
            'nama_ilap': ilap_obj.nama_ilap if ilap_obj else '-',
            'sub_jenis_data': sub.nama_sub_jenis_data if sub else '-',
            'jenis_data': sub.nama_jenis_data if sub else '-',
            'periode_data': _format_periode_tiket(t),
            'status_data': status_data,
            'status_penelitian': status_penelitian,
            'jumlah_baris_diterima': format_number_with_separator(t.baris_diterima),
            'jumlah_data_diterima': format_number_with_separator(t.baris_diterima),
            'jumlah_baris_lengkap': format_number_with_separator(t.baris_lengkap) if t.baris_lengkap is not None else '-',
            'jumlah_baris_tidak_lengkap': format_number_with_separator(t.baris_tidak_lengkap) if t.baris_tidak_lengkap is not None else '-',
            'dasar_hukum': dasar_hukum if dasar_hukum else '-',
            'nomor_surat_pengantar': t.nomor_surat_pengantar or '-',
            'tanggal_surat_pengantar': _format_date_indonesian(t.tanggal_surat_pengantar),
            'tanggal_terima_dip': _format_date_indonesian(t.tgl_terima_dip),
            'tanggal_kirim_pide': _format_date_indonesian(t.tgl_kirim_pide),
        })

    now_ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    tahun_data_list = sorted({str(t.tahun) for t in selected_tickets if t.tahun})
    nomor_tanda_terima_text = ', '.join(nomor_tanda_terima_list) if nomor_tanda_terima_list else '-'

    template_vars = {
        '{{nomor_tiket}}': nomor_tanda_terima_text,
        '{{nomor_tanda_terima}}': nomor_tanda_terima_text,
        '{{tanggal_tanda_terima}}': ', '.join(tanggal_tanda_terima_list) if tanggal_tanda_terima_list else '-',
        '{{tahun_data}}': ', '.join(tahun_data_list) if tahun_data_list else '-',
        '{{diterima_dari}}': diterima_dari,
        '{{nama_kantor}}': diterima_dari,
        '{{nomor_surat_pengantar}}': ', '.join(nomor_surat_list) if nomor_surat_list else '-',
        '{{tanggal_surat_pengantar}}': ', '.join(tanggal_surat_list) if tanggal_surat_list else '-',
        '{{tanggal_penerimaan}}': ', '.join(tanggal_surat_list) if tanggal_surat_list else '-',
        '{{nama_ilap}}': ilap.nama_ilap if ilap else '-',
        '{{jenis_data}}': 'Terlampir',
        '{{periode_data}}': ', '.join(periode_list) if periode_list else '-',
        '{{bentuk_data}}': ', '.join(bentuk_data_list) if bentuk_data_list else '-',
        '{{tanggal_terima_dip}}': _format_date_indonesian(first_ticket.tgl_terima_dip),
        '{{cara_penyampaian}}': ', '.join(cara_penyampaian_list) if cara_penyampaian_list else '-',
        '{{nama_pic_p3de}}': '-',
        '{{nama_pic}}': '-',
        '{{email_pic}}': '-',
        '{{telepon_pic}}': '-',
        '{{nama_tabel}}': 'Terlampir',
        '{{jumlah_record}}': '-',
        '{{ukuran_file}}': '-',
        '{{tanggal_cetak}}': _format_date_indonesian(datetime.now()),
        '{{jumlah_tiket}}': str(len(selected_tickets)),
        '{{jenis_dokumen}}': doc_type,
    }

    template = DocxTemplate.objects.filter(jenis_dokumen=template_jenis, active=True).first() if template_jenis else None
    if template and template.file_template:
        try:
            doc_buffer = fill_template_with_data(
                template.file_template.open('rb'),
                template_vars,
                row_data=row_data,
            )
            response = HttpResponse(
                doc_buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{title_prefix}_{now_ts}.docx"'
            return response
        except Exception:
            pass

    # Fallback document if template unavailable/error.
    headers = ['No', 'Nomor Tiket', 'Nama ILAP', 'Sub Jenis Data', 'Periode Data', 'Baris Diterima', 'Status Penelitian']
    rows = [
        [
            r['nomor'],
            r['nomor_tiket'],
            r['nama_ilap'],
            r['sub_jenis_data'],
            r['periode_data'],
            r['jumlah_baris_diterima'],
            r['status_penelitian'],
        ]
        for r in row_data
    ]
    doc = _build_table_doc(f'{title_prefix} ({len(rows)} tiket)', headers, rows)
    buffer = BytesIO()
    doc.save(buffer)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{title_prefix}_{now_ts}.docx"'
    return response


@login_required
@user_passes_test(_is_p3de_user)
@require_http_methods(['GET', 'POST'])
def bulk_pkdi_klarifikasi(request):
    # Restrict ILAP list to user's active P3DE assignments unless admin
    if request.user.is_superuser or request.user.groups.filter(name__in=['admin', 'admin_p3de']).exists():
        ilap_options = ILAP.objects.order_by('nama_ilap')
    else:
        ilap_ids = get_active_p3de_ilap_ids(request.user)
        ilap_options = ILAP.objects.filter(id__in=ilap_ids).order_by('nama_ilap')

    ilap_id = request.GET.get('ilap_id', '')
    tanggal_terima = request.GET.get('tanggal_terima', '')
    doc_type = request.GET.get('doc_type', 'pkdi_lengkap')

    tickets = []
    if tanggal_terima and doc_type in ['pkdi_lengkap', 'pkdi_sebagian', 'klarifikasi']:
        tanggal_obj = _parse_date(tanggal_terima)
        if tanggal_obj:
            # If ilap_id is empty or 'semua', show all tickets for the date
            if not ilap_id or ilap_id == 'semua':
                qs = Tiket.objects.filter(
                    tgl_terima_dip__date=tanggal_obj,
                    tanda_terima=True,
                ).select_related(
                    'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                    'id_periode_data__id_periode_pengiriman',
                    'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                    'id_status_penelitian',
                ).prefetch_related(
                    'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
                ).order_by('id')
                tickets = list(_apply_doc_type_filter(qs, doc_type))
            else:
                qs = _base_queryset(ilap_id, tanggal_obj)
                tickets = list(_apply_doc_type_filter(qs, doc_type))

    if request.method == 'POST':
        ilap_id = request.POST.get('ilap_id', '')
        tanggal_terima = request.POST.get('tanggal_terima', '')
        doc_type = request.POST.get('doc_type', '')
        selected_ids = request.POST.getlist('ticket_ids')

        tanggal_obj = _parse_date(tanggal_terima)
        if not tanggal_obj or doc_type not in ['pkdi_lengkap', 'pkdi_sebagian', 'klarifikasi']:
            messages.error(request, 'Parameter filter tidak valid.')
            return redirect('bulk_pkdi_klarifikasi')

        # Handle 'semua' or empty ilap_id
        if not ilap_id or ilap_id == 'semua':
            base_qs = Tiket.objects.filter(
                tgl_terima_dip__date=tanggal_obj,
                tanda_terima=True,
            ).select_related(
                'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                'id_periode_data__id_periode_pengiriman',
                'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                'id_status_penelitian',
            ).prefetch_related(
                'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
            ).order_by('id')
            base_qs = _apply_doc_type_filter(base_qs, doc_type)
        else:
            base_qs = _apply_doc_type_filter(_base_queryset(ilap_id, tanggal_obj), doc_type)
        
        selected_tickets = list(base_qs.filter(id__in=selected_ids).order_by('id'))

        if not selected_tickets:
            messages.warning(request, 'Pilih minimal 1 tiket untuk digenerate.')
            query = f'?ilap_id={ilap_id}&tanggal_terima={tanggal_terima}&doc_type={doc_type}'
            return redirect(f"/bulk-generate/pkdi-klarifikasi/{query}")

        return _generate_docx_for_tickets(selected_tickets, doc_type, 'bulk_pkdi_klarifikasi')

    return render(request, 'bulk_documents/pkdi_klarifikasi.html', {
        'page_title': 'Generate PKDI / Klarifikasi (Bulk)',
        'ilap_options': ilap_options,
        'tickets': tickets,
        'selected_ilap_id': str(ilap_id),
        'selected_tanggal_terima': tanggal_terima,
        'selected_doc_type': doc_type,
    })


@login_required
@user_passes_test(_is_p3de_user)
@require_http_methods(['GET', 'POST'])
def bulk_nd_pengantar_pide(request):
    # Restrict ILAP list to user's active P3DE assignments unless admin
    if request.user.is_superuser or request.user.groups.filter(name__in=['admin', 'admin_p3de']).exists():
        ilap_options = ILAP.objects.order_by('nama_ilap')
    else:
        ilap_ids = get_active_p3de_ilap_ids(request.user)
        ilap_options = ILAP.objects.filter(id__in=ilap_ids).order_by('nama_ilap')

    ilap_id = request.GET.get('ilap_id', '')
    tanggal_kirim_pide = request.GET.get('tanggal_kirim_pide', '')
    tickets = []

    if tanggal_kirim_pide:
        tanggal_obj = _parse_date(tanggal_kirim_pide)
        if tanggal_obj:
            # If ilap_id is empty or 'semua', show all tickets for the date
            if not ilap_id or ilap_id == 'semua':
                tickets = list(
                    Tiket.objects.filter(
                        tgl_kirim_pide__date=tanggal_obj,
                        tanda_terima=True,
                    ).select_related(
                        'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                        'id_periode_data__id_periode_pengiriman',
                        'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                        'id_status_penelitian',
                    ).prefetch_related(
                        'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
                    )
                    .order_by('id')
                )
            else:
                tickets = list(
                    Tiket.objects.filter(
                        id_periode_data__id_sub_jenis_data_ilap__id_ilap_id=ilap_id,
                        tgl_kirim_pide__date=tanggal_obj,
                        tanda_terima=True,
                    ).select_related(
                        'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                        'id_periode_data__id_periode_pengiriman',
                        'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                        'id_status_penelitian',
                    ).prefetch_related(
                        'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
                    )
                    .order_by('id')
                )

    if request.method == 'POST':
        ilap_id = request.POST.get('ilap_id', '')
        tanggal_kirim_pide = request.POST.get('tanggal_kirim_pide', '')
        selected_ids = request.POST.getlist('ticket_ids')

        tanggal_obj = _parse_date(tanggal_kirim_pide)
        if not tanggal_obj:
            messages.error(request, 'Parameter filter tidak valid.')
            return redirect('bulk_nd_pengantar_pide')

        # Handle 'semua' or empty ilap_id
        if not ilap_id or ilap_id == 'semua':
            base_qs = Tiket.objects.filter(
                tgl_kirim_pide__date=tanggal_obj,
                tanda_terima=True,
            ).select_related(
                'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                'id_periode_data__id_periode_pengiriman',
                'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                'id_status_penelitian',
            ).prefetch_related(
                'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
            ).order_by('id')
        else:
            base_qs = Tiket.objects.filter(
                id_periode_data__id_sub_jenis_data_ilap__id_ilap_id=ilap_id,
                tgl_kirim_pide__date=tanggal_obj,
                tanda_terima=True,
            ).select_related(
                'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                'id_periode_data__id_periode_pengiriman',
                'id_periode_data__id_sub_jenis_data_ilap__id_status_data',
                'id_status_penelitian',
            ).prefetch_related(
                'id_periode_data__id_sub_jenis_data_ilap__klasifikasijenisdata_set__id_klasifikasi_tabel',
            ).order_by('id')
        
        selected_tickets = list(base_qs.filter(id__in=selected_ids).order_by('id'))

        if not selected_tickets:
            messages.warning(request, 'Pilih minimal 1 tiket untuk digenerate.')
            query = f'?ilap_id={ilap_id}&tanggal_kirim_pide={tanggal_kirim_pide}'
            return redirect(f"/bulk-generate/nd-pengantar-pide/{query}")

        return _generate_docx_for_tickets(selected_tickets, 'nd_pengantar', 'bulk_nd_pengantar_pide')

    return render(request, 'bulk_documents/nd_pengantar_pide.html', {
        'page_title': 'Generate ND Pengantar PIDE (Bulk)',
        'ilap_options': ilap_options,
        'tickets': tickets,
        'selected_ilap_id': str(ilap_id),
        'selected_tanggal_kirim_pide': tanggal_kirim_pide,
    })
