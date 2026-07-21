"""Tests covering remaining gaps in diamond_web/views/tiket/documents.py.

Uncovered lines:
- 106-107: ImportError handler for `from docx import Document`
- 156: `dasar_hukum_map.setdefault(...)` body of for-loop (needs KlasifikasiJenisData record)
- 291-300: `status_penelitian` branches (Lengkap / Sebagian / Tidak) inside lampiran row_data loop
- 324-333: pkdi_lengkap / pkdi_sebagian / klarifikasi / nd_pengantar / else filename branches
            when an active DocxTemplate is found (template path)
"""
import sys
import pytest
from io import BytesIO
from unittest.mock import patch
from django.urls import reverse
from django.contrib.auth.models import Group
from django.core.files.uploadedfile import SimpleUploadedFile
from django.utils import timezone

from diamond_web.models import TiketPIC, TandaTerimaData
from diamond_web.models.detil_tanda_terima import DetilTandaTerima
from diamond_web.tests.conftest import (
    TiketFactory, TiketPICFactory, UserFactory, ILAPFactory,
    KategoriWilayahFactory, KanwilFactory, KPPFactory,
    PeriodePengirimanFactory, PeriodeJenisDataFactory, JenisDataILAPFactory,
    DocxTemplateFactory, BentukDataFactory, CaraPenyampaianFactory,
    KlasifikasiJenisDataFactory,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes(paragraphs=None, table_rows=None):
    """Create a minimal DOCX as BytesIO using python-docx."""
    from docx import Document
    doc = Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if table_rows:
        ncols = max(len(r) for r in table_rows) if table_rows else 1
        table = doc.add_table(rows=len(table_rows), cols=ncols)
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                table.cell(i, j).paragraphs[0].add_run(cell_text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_superuser():
    user = UserFactory()
    user.is_superuser = True
    user.save()
    return user


def _make_p3de_user():
    user = UserFactory()
    grp, _ = Group.objects.get_or_create(name='user_p3de')
    user.groups.add(grp)
    return user


def _make_tiket_with_tanda_terima_and_jenis_data():
    """Create a tiket with tanda_terima, TandaTerima, DetilTandaTerima and JenisDataILAP chain."""
    jd = JenisDataILAPFactory()
    pp = PeriodePengirimanFactory()
    pd = PeriodeJenisDataFactory(id_sub_jenis_data_ilap=jd, id_periode_pengiriman=pp)
    tiket = TiketFactory(
        status_tiket=1,
        tanda_terima=True,
        id_periode_data=pd,
        id_bentuk_data=BentukDataFactory(),
        id_cara_penyampaian=CaraPenyampaianFactory(),
    )
    return tiket, jd


def _make_tiket_with_tt_and_detil():
    """Create tiket with TandaTerima + DetilTandaTerima for template tests."""
    user = _make_superuser()
    jd = JenisDataILAPFactory()
    pp = PeriodePengirimanFactory()
    pd = PeriodeJenisDataFactory(id_sub_jenis_data_ilap=jd, id_periode_pengiriman=pp)
    tiket = TiketFactory(
        status_tiket=1,
        tanda_terima=True,
        id_periode_data=pd,
        id_bentuk_data=BentukDataFactory(),
        id_cara_penyampaian=CaraPenyampaianFactory(),
    )
    ilap = jd.id_ilap
    tt = TandaTerimaData.objects.create(
        nomor_tanda_terima=88880,
        tahun_terima=2080,
        tanggal_tanda_terima=timezone.now(),
        id_ilap=ilap,
        id_perekam=user,
        active=True,
    )
    DetilTandaTerima.objects.create(id_tanda_terima=tt, id_tiket=tiket)
    return tiket, user


def _get_region_type_jenis_dokumen(tiket, doc_type):
    """Compute the jenis_dokumen key used by the view for template lookup."""
    try:
        ilap = tiket.id_periode_data.id_sub_jenis_data_ilap.id_ilap
        if ilap and ilap.id_kategori_wilayah:
            desc = ilap.id_kategori_wilayah.deskripsi.lower()
            region = 'regional' if 'regional' in desc else 'nasional_internasional'
        else:
            region = 'regional'
    except Exception:
        region = 'regional'

    doc_type_map = {
        'tanda_terima': f'tanda_terima_{region}',
        'lampiran': f'lampiran_tanda_terima_{region}',
        'register': 'register_penerimaan_data',
        'pkdi_lengkap': f'surat_pkdi_{region}_lengkap',
        'pkdi_sebagian': f'surat_pkdi_{region}_sebagian',
        'klarifikasi': 'surat_klarifikasi',
        'nd_pengantar': 'nd_pengantar_pide',
    }
    return doc_type_map.get(doc_type, f'tanda_terima_{region}')


def _make_template_for(tiket, doc_type):
    """Create an active DocxTemplate whose jenis_dokumen matches the view's lookup key."""
    docx_bytes = _make_docx_bytes(
        paragraphs=['{{nomor_tanda_terima}}'],
        table_rows=[['Header'], ['{{row.nama_ilap}}']],
    ).getvalue()
    jenis = _get_region_type_jenis_dokumen(tiket, doc_type)
    return DocxTemplateFactory(
        jenis_dokumen=jenis,
        active=True,
        file_template=SimpleUploadedFile(f'tmpl_{doc_type}.docx', docx_bytes),
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestDocumentsRemainingGaps:
    """Cover the remaining uncovered lines in documents.py."""

    # -----------------------------------------------------------------------
    # Lines 106-107: ImportError handler
    # -----------------------------------------------------------------------
    def test_import_error_lines_106_107(self, client, db):
        """When `from docx import Document` raises ImportError, return HTTP 500.

        Strategy: patch builtins.__import__ to raise ImportError only when
        importing 'docx'. Since diamond_web.utils.docx_template is already
        loaded (imported at startup), its existing `Document` reference is
        unaffected. Only the LOCAL `from docx import Document` inside the
        view function will fail, exercising the except-ImportError branch.
        """
        import builtins

        user = _make_superuser()
        client.force_login(user)

        tiket = TiketFactory(tanda_terima=True)
        TiketPICFactory(id_tiket=tiket, id_user=user, role=TiketPIC.Role.P3DE, active=True)

        # Ensure utils.docx_template is already imported before patching
        import diamond_web.utils.docx_template  # noqa: F401

        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == 'docx':
                raise ImportError('Mocked: python-docx not installed')
            return real_import(name, *args, **kwargs)

        with patch.object(builtins, '__import__', side_effect=fake_import):
            resp = client.get(
                reverse('tiket_documents_download', args=[tiket.pk]),
                {'doc_type': 'tanda_terima'},
            )
        assert resp.status_code == 500
        assert b'python-docx' in resp.content

    # -----------------------------------------------------------------------
    # Line 156: dasar_hukum_map body (inside for-loop over KlasifikasiJenisData)
    # -----------------------------------------------------------------------
    def test_dasar_hukum_map_setdefault_line_156(self, client, db):
        """Line 156: loop body executes when KlasifikasiJenisData exists for the tiket's jenis data."""
        user = _make_superuser()
        client.force_login(user)

        tiket, jd = _make_tiket_with_tanda_terima_and_jenis_data()

        # Create KlasifikasiJenisData linking to the tiket's jenis_data → triggers line 156
        KlasifikasiJenisDataFactory(id_sub_jenis_data=jd)

        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'tanda_terima'},
        )
        assert resp.status_code == 200

    # -----------------------------------------------------------------------
    # Lines 291-300: status_penelitian branches inside lampiran row_data loop
    # Strategy: patch get_object_or_404 to inject status_penelitian on the tiket
    # then request lampiran doc_type with an active template
    # (tiket_rows=[tiket] since no DetilTandaTerima exists)
    # -----------------------------------------------------------------------
    def _setup_lampiran_test(self, client, status_nilai):
        """Helper: sets up user, tiket, lampiran template and patches get_object_or_404."""
        user = _make_superuser()
        client.force_login(user)
        tiket, jd = _make_tiket_with_tanda_terima_and_jenis_data()
        _make_template_for(tiket, 'lampiran')

        from django.shortcuts import get_object_or_404 as real_g404

        def mock_g404(qs, **kwargs):
            obj = real_g404(qs, **kwargs)
            obj.status_penelitian = status_nilai
            return obj

        return tiket, mock_g404

    def test_status_penelitian_lengkap_lines_291_293(self, client, db):
        """Lines 291-293: status_penelitian contains 'lengkap' (but not sebagian/tidak)."""
        tiket, mock_g404 = self._setup_lampiran_test(client, 'lengkap')
        with patch('diamond_web.views.tiket.documents.get_object_or_404', side_effect=mock_g404):
            resp = client.get(
                reverse('tiket_documents_download', args=[tiket.pk]),
                {'doc_type': 'lampiran'},
            )
        assert resp.status_code == 200

    def test_status_penelitian_sebagian_lines_294_297(self, client, db):
        """Lines 294-297: status_penelitian contains 'sebagian' (but not 'lengkap')."""
        tiket, mock_g404 = self._setup_lampiran_test(client, 'sebagian saja')
        with patch('diamond_web.views.tiket.documents.get_object_or_404', side_effect=mock_g404):
            resp = client.get(
                reverse('tiket_documents_download', args=[tiket.pk]),
                {'doc_type': 'lampiran'},
            )
        assert resp.status_code == 200

    def test_status_penelitian_tidak_lines_298_300(self, client, db):
        """Lines 298-300: status_penelitian contains 'tidak' (but not 'lengkap' or 'sebagian')."""
        tiket, mock_g404 = self._setup_lampiran_test(client, 'tidak tersedia')
        with patch('diamond_web.views.tiket.documents.get_object_or_404', side_effect=mock_g404):
            resp = client.get(
                reverse('tiket_documents_download', args=[tiket.pk]),
                {'doc_type': 'lampiran'},
            )
        assert resp.status_code == 200

    # -----------------------------------------------------------------------
    # Lines 324-333: filename branches for pkdi/klarifikasi/nd_pengantar/else
    #                when an active template IS found (template path)
    # -----------------------------------------------------------------------
    def test_pkdi_lengkap_filename_line_324_325(self, client, db):
        """Lines 324-325: pkdi_lengkap doc_type with active template sets correct filename."""
        tiket, user = _make_tiket_with_tt_and_detil()
        client.force_login(user)
        _make_template_for(tiket, 'pkdi_lengkap')
        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'pkdi_lengkap'},
        )
        assert resp.status_code == 200
        assert 'surat_pkdi_lengkap' in resp.get('Content-Disposition', '')

    def test_pkdi_sebagian_filename_line_326_327(self, client, db):
        """Lines 326-327: pkdi_sebagian doc_type with active template."""
        tiket, user = _make_tiket_with_tt_and_detil()
        client.force_login(user)
        _make_template_for(tiket, 'pkdi_sebagian')
        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'pkdi_sebagian'},
        )
        assert resp.status_code == 200
        assert 'surat_pkdi' in resp.get('Content-Disposition', '')

    def test_klarifikasi_filename_line_328_329(self, client, db):
        """Lines 328-329: klarifikasi doc_type with active template."""
        tiket, user = _make_tiket_with_tt_and_detil()
        client.force_login(user)
        _make_template_for(tiket, 'klarifikasi')
        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'klarifikasi'},
        )
        assert resp.status_code == 200
        assert 'surat_klarifikasi' in resp.get('Content-Disposition', '')

    def test_nd_pengantar_filename_line_330_331(self, client, db):
        """Lines 330-331: nd_pengantar doc_type with active template."""
        tiket, user = _make_tiket_with_tt_and_detil()
        client.force_login(user)
        _make_template_for(tiket, 'nd_pengantar')
        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'nd_pengantar'},
        )
        assert resp.status_code == 200
        assert 'nd_pengantar_pide' in resp.get('Content-Disposition', '')

    def test_else_filename_line_332_333(self, client, db):
        """Lines 332-333: else branch (tanda_terima or unknown doc_type) with active template."""
        tiket, user = _make_tiket_with_tt_and_detil()
        client.force_login(user)
        _make_template_for(tiket, 'tanda_terima')
        resp = client.get(
            reverse('tiket_documents_download', args=[tiket.pk]),
            {'doc_type': 'tanda_terima'},
        )
        assert resp.status_code == 200
        assert 'tanda_terima' in resp.get('Content-Disposition', '')
